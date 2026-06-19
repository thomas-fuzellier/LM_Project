from docx import Document
from datetime import datetime
from core.profil import charger_profil


def remplacer_dans_paragraphe(paragraph, mapping):
    texte_complet = "".join(run.text for run in paragraph.runs)
    if not any(k in texte_complet for k in mapping):
        return
    for k, v in mapping.items():
        texte_complet = texte_complet.replace(k, v)
    if paragraph.runs:
        paragraph.runs[0].text = texte_complet
        for run in paragraph.runs[1:]:
            run.text = ""


def remplacer_placeholder(doc, mapping):
    for p in doc.paragraphs:
        remplacer_dans_paragraphe(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    remplacer_dans_paragraphe(p, mapping)
    for section in doc.sections:
        for p in section.header.paragraphs:
            remplacer_dans_paragraphe(p, mapping)
        for p in section.footer.paragraphs:
            remplacer_dans_paragraphe(p, mapping)


def generer_lettre(template_path, output_path, entreprise, poste):
    doc = Document(template_path)
    donnees = charger_profil()

    mapping = {
        "{entreprise}": entreprise,
        "{poste}": poste,
        "{date}": datetime.now().strftime("%d/%m/%Y"),
        "{nom}": donnees.get("nom", ""),
        "{prenom}": donnees.get("prenom", ""),
        "{email}": donnees.get("email", ""),
        "{telephone}": donnees.get("telephone", ""),
        "{ville}": donnees.get("ville", ""),
    }

    remplacer_placeholder(doc, mapping)
    doc.save(output_path)