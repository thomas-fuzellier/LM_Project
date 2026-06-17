import tkinter as tk
from datetime import datetime
from docx import Document
import win32com.client as win32
import os


# ---------------- WORD ---------------- #

def remplacer_placeholder(doc, mapping):
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                p.text = p.text.replace(k, v)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k, v in mapping.items():
                    if k in cell.text:
                        cell.text = cell.text.replace(k, v)


def generer_lettre(template_path, output_docx, entreprise, poste):
    doc = Document(template_path)

    mapping = {
        "{entreprise}": entreprise,
        "{poste}": poste,
        "{date}": datetime.now().strftime("%d/%m/%Y")
    }

    remplacer_placeholder(doc, mapping)
    doc.save(output_docx)


# ---------------- PDF (Word COM) ---------------- #
def docx_to_pdf(docx_path, pdf_path):
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    # >>> OPTIMISATIONS ICI <<<
    word.ScreenUpdating = False
    word.Options.SaveInterval = 0

    doc = word.Documents.Open(docx_path, ReadOnly=1)

    doc.SaveAs(pdf_path, FileFormat=17)

    doc.Close(False)
    word.Quit()


# ---------------- APP ---------------- #

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Générateur de lettres")
        self.root.geometry("420x320")

        self.type_job = "autre"

        # --- INPUTS ---
        tk.Label(root, text="Nom de l'entreprise").pack()
        self.entry_entreprise = tk.Entry(root, width=40)
        self.entry_entreprise.pack()

        tk.Label(root, text="Nom du poste").pack()
        self.entry_poste = tk.Entry(root, width=40)
        self.entry_poste.pack()

        # --- TOGGLE BUTTON ---
        self.btn_type = tk.Button(
            root,
            text="Secteur d'activité : autre",
            command=self.toggle_type,
            width=20
        )
        self.btn_type.pack(pady=10)

        # --- GENERATE BUTTON ---
        tk.Button(
            root,
            text="Générer PDF",
            command=self.generer,
            width=20
        ).pack(pady=10)

        # --- STATUS ---
        self.label_status = tk.Label(root, text="")
        self.label_status.pack()

    # ---------------- LOGIQUE ---------------- #

    def toggle_type(self):
        if self.type_job == "defense":
            self.type_job = "autre"
            self.btn_type.config(text="Secteur d'activité : autre")
        else:
            self.type_job = "defense"
            self.btn_type.config(text="Secteur d'activité : defense")

    def generer(self):
        entreprise = self.entry_entreprise.get().strip()
        poste = self.entry_poste.get().strip()

        base_path = r"C:\Users\Utilisateur\Desktop\LM_Project"
        output_docx = os.path.join(base_path, "lettre_generee.docx")
        output_pdf = os.path.join(base_path, "lettre_generee.pdf")

        if self.type_job == "defense":
            template = os.path.join(base_path, "lettre_template_defense.docx")
        else:
            template = os.path.join(base_path, "lettre_template.docx")

        # génération docx
        generer_lettre(template, output_docx, entreprise, poste)

        # conversion pdf
        docx_to_pdf(output_docx, output_pdf)

        # optionnel : suppression docx
        if os.path.exists(output_docx):
            os.remove(output_docx)

        self.label_status.config(text="PDF généré ✔")


# ---------------- MAIN ---------------- #

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()